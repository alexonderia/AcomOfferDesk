"""Fill column 'Atomic roles in Keycloak' in the project responsibility matrix docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

DOCX = Path(r"c:\Users\alexonderia\Desktop\доки по проекту\матрица ответственности проекта.docx")

# User-facing function (column «Функция») -> atomic client roles in Keycloak (acom-api).
FUNCTION_TO_ATOMIC_ROLES: dict[str, str] = {
    "Создание сотрудника": "users.create",
    "Просмотр списка всех пользователей": "users.read",
    "Просмотр списка подчиненных (экономистов)": "users.read",
    "Кандидаты в руководители": "users.read",
    "Профиль текущего пользователя": "profile.manage_own",
    "Смена логина и пароля": "users.login.update, users.password.update",
    "Обновление собственного профиля": "profile.manage_own",
    "Обновление своих контактов компании": "company_contacts.manage_own",
    "Период недоступности для себя": "unavailability.manage_own",
    "Профиль подчиненного": "profile.manage_any, users.read",
    "Период недоступности подчиненного": "unavailability.manage_subordinate",
    "Справочник экономистов для ответственного по заявке": (
        "— (role gate: app.superadmin, app.project_manager, app.lead_economist; отдельного permission нет)"
    ),
    "Справочник контрагентов для создания КП": "requests.create",
    "Создание контрагента вручную (manual)": "contractors.manual.create",
    "Изменение данных manual-контрагента": "contractors.manual.manage",
    "Обновление статуса всех пользователей": "users.status.update",
    "Обновление статуса подчиненных": "users.status.update",
    "Обновление роли любого пользователя": "users.role.update_any",
    "Обновление роли подчиненных": "users.role.update_economy",
    "Смена руководителя любому экономисту": "users.manager.update",
    "Смена руководителя подчиненному": "users.manager.update",
    "Список заявок (staff-view)": "requests.read",
    "Список открытых заявок": "requests.open.read",
    "Список заявок с офферами контрагента": "requests.offered.read",
    "Детали заявки + офферы": (
        "requests.read, offers.contractor_info.read (список КП; без offers.workspace.read)"
    ),
    "Создание заявки": "requests.create",
    "Обновление заявки (цены/статус/дедлайн/план)": (
        "requests.owner.change; requests.update, requests.pricing.update, "
        "requests.deadline.update, requests.status.update"
    ),
    "Ручная email-рассылка по заявке": "requests.email_notifications.send",
    "Добавление файла в заявку": "requests.files.upload",
    "Удаление файла из заявки": "requests.files.delete",
    "Пометка alert по удаленным офферам как просмотренного": "requests.deleted_alerts.mark_viewed",
    "Скачивание файла": "files.download",
    "Карточка контрагента": "offers.contractor_info.read",
    "Contractor-view заявки": "requests.contractor_view.read",
    "Создание оффера контрагентом": "offers.create",
    "Создание manual-оффера сотрудником": "offers.manual.create",
    "Рабочее пространство оффера": "offers.workspace.read",
    "Смена статуса оффера": "offers.status.update",
    "Изменение суммы оффера": "offers.amount.update",
    "Загрузка файла в оффер": "offers.files.upload",
    "Удаление файла из оффера": "offers.files.delete",
    "Список сообщений чата": "chat.read",
    "Отправка текстового сообщения": "chat.message.send",
    "Сообщение с вложениями": "chat.message.attach",
    "Подтверждение доставки сообщений": "chat.receipts.mark_received",
    "Подтверждение прочтения сообщений": "chat.receipts.mark_read",
    "Подписка/синхронизация чата": "chat.read",
    "План-дашборд (месяц/диапазон)": "dashboard.plans.read",
    "Статистика заявок по плану": "dashboard.plans.read",
    "Дерево планов пользователя": "dashboard.plans.read",
    "Моя сводка по плану": "dashboard.plans.read",
    "Кандидаты делегирования плана": "dashboard.plans.read",
    "Опции планов (селекторы)": "dashboard.plans.read",
    "Создание root-плана": "dashboard.plans.read",
    "Создание подплана": "dashboard.plans.read",
    "Делегирование плана": "dashboard.plans.read",
    "Редактирование плана": "dashboard.plans.read",
    "Удаление дочернего плана": "dashboard.plans.read",
    "Закрытие плана": "dashboard.plans.read",
    "Создание обратной связи": "feedback.create",
    "Просмотр списка обратной связи": "feedback.read",
    "Доступ к нормативным документам в заявочном контуре": "normative_files.read",
    "Загрузка нормативного файла": "normative_files.create, normative_files.manage",
}


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    document = Document(DOCX)
    if not document.tables:
        raise RuntimeError("No tables found in document")

    table = document.tables[0]
    updated = 0
    missing: list[str] = []

    for row in table.rows[1:]:
        function = row.cells[1].text.strip()
        if not function:
            continue
        atomic = FUNCTION_TO_ATOMIC_ROLES.get(function)
        if atomic is None:
            missing.append(function)
            continue
        if row.cells[0].text.strip() != atomic:
            row.cells[0].text = atomic
            updated += 1

    document.save(DOCX)
    print(f"Saved: {DOCX}")
    print(f"Updated rows: {updated}")
    if missing:
        print("Unmapped functions:")
        for item in missing:
            print(f"  - {item}")


if __name__ == "__main__":
    main()
