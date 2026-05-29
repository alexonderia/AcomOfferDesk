# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCX = Path(r"c:\Users\alexonderia\Downloads\Reorganized-user-guide.docx")
OUT = Path(__file__).resolve().parents[1] / "docs/reports/docx-structure-sample.json"


def iter_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("p", Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            yield ("tbl", Table(child, document))


def runs_to_markdown(paragraph: Paragraph) -> str:
    chunks: list[str] = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            chunks.append(f"**{text}**")
        else:
            chunks.append(text)
    return "".join(chunks).strip()


def has_numbering(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    return p_pr.numPr is not None


def table_to_dict(table: Table) -> dict:
    rows = [[runs_to_markdown(cell.paragraphs[0]) if cell.paragraphs else cell.text.strip() for cell in row.cells] for row in table.rows]
    # merge duplicate cells in merged rows - docx repeats cell text
    headers = rows[0] if rows else []
    body_rows = rows[1:] if len(rows) > 1 else []
    return {"headers": headers, "rows": body_rows}


def main() -> None:
    doc = Document(DOCX)
    sample: list[dict] = []
    count = 0
    for kind, block in iter_blocks(doc):
        if count > 80:
            break
        if kind == "p":
            p: Paragraph = block
            text = p.text.strip()
            if not text:
                continue
            sample.append(
                {
                    "kind": "p",
                    "style": p.style.name if p.style else "",
                    "numbered": has_numbering(p),
                    "text": runs_to_markdown(p),
                }
            )
        else:
            sample.append({"kind": "tbl", **table_to_dict(block)})
        count += 1

    OUT.write_text(json.dumps(sample, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
